"""
Generate dokumentasi DOCX untuk setup database PostgreSQL ALTIVEX
+ contoh log MQTT publish + insert + WebSocket broadcast.

Output: docs/DOKUMENTASI_DATABASE_DAN_LOGS.docx

Usage:
    pip3 install python-docx
    python3 scripts/generate-db-docs.py
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------
# Style helpers
# --------------------------------------------------------------------
def shade_cell(cell, hex_color):
    """Shade a table cell background with a hex color."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_code_block(doc, text, lang_label=None):
    """Add a monospace block paragraph with grey background."""
    if lang_label:
        p = doc.add_paragraph()
        run = p.add_run(f'[{lang_label}]')
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)

    # Background fill via paragraph property
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    p_pr.append(shd)

    # Border (light grey)
    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        pBdr.append(b)
    p_pr.append(pBdr)

    run = p.add_run(text)
    run.font.name = 'Consolas'
    # Set font for east asian + ascii (Word quirk)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    run.font.size = Pt(9)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_bold_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    return p


# --------------------------------------------------------------------
# Build document
# --------------------------------------------------------------------
def build_document(out_path):
    doc = Document()

    # --- Page setup --------------------------------------------------
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ================================================================
    # Cover / Title
    # ================================================================
    title = doc.add_heading('Dokumentasi Setup Database\ndan Pengamatan Log Sistem', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run('Sistem Pelacak Pendaki ALTIVEX')
    run.italic = True
    run.font.size = Pt(14)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f'Tanggal disusun: {datetime.now().strftime("%d %B %Y")}')
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()
    doc.add_paragraph()

    # ================================================================
    # Pendahuluan
    # ================================================================
    add_heading(doc, '1. Pendahuluan', level=1)
    doc.add_paragraph(
        'Dokumen ini menyajikan dokumentasi teknis setup basis data PostgreSQL '
        'pada sistem ALTIVEX serta contoh nyata log keluaran sistem ketika '
        'menerima data telemetri dari device pendaki via MQTT. Dokumen ini '
        'ditujukan sebagai bukti operasional bahwa pipeline data berjalan '
        'utuh dari pengirim (ESP32 + GPS NEO-6M) hingga tersimpan di basis '
        'data dan disebarkan ke dashboard penjaga pos secara real-time.'
    )
    doc.add_paragraph(
        'Ruang lingkup dokumen mencakup tiga bagian utama: (1) skema basis '
        'data dan migrasi otomatis yang dijalankan saat backend pertama kali '
        'dijalankan, (2) konfigurasi broker MQTT Mosquitto dengan otentikasi '
        'password, dan (3) contoh log keluaran backend Rust dan basis data '
        'saat menerima publish dari device.'
    )

    # ================================================================
    # Skema Basis Data
    # ================================================================
    add_heading(doc, '2. Skema Basis Data', level=1)
    doc.add_paragraph(
        'Sistem ALTIVEX menggunakan PostgreSQL 15 sebagai sistem manajemen '
        'basis data utama. Skema terdiri atas dua tabel inti: log_sensor '
        'untuk menyimpan setiap publish koordinat pendaki, dan pendaki '
        'untuk menyimpan registrasi pendaki yang sedang aktif maupun yang '
        'telah selesai mendaki. Migrasi skema dijalankan secara otomatis '
        'oleh backend saat startup melalui perintah CREATE TABLE IF NOT '
        'EXISTS sehingga deployment baru tidak memerlukan langkah migrasi '
        'manual.'
    )

    add_heading(doc, '2.1 Tabel log_sensor', level=2)
    doc.add_paragraph(
        'Tabel log_sensor menampung setiap titik koordinat pendaki yang '
        'diterima backend, baik melalui MQTT, HTTP REST, maupun Serial '
        'bridge. Skema dibuat dengan kolom timestamp ber-default '
        'CURRENT_TIMESTAMP dan kolom battery sebagai SMALLINT optional.'
    )

    add_code_block(doc,
        'CREATE TABLE IF NOT EXISTS log_sensor (\n'
        '    id            SERIAL       PRIMARY KEY,\n'
        '    id_perangkat  VARCHAR(50)  NOT NULL,\n'
        '    latitude      DOUBLE PRECISION NOT NULL,\n'
        '    longitude     DOUBLE PRECISION NOT NULL,\n'
        '    battery       SMALLINT,\n'
        '    timestamp     TIMESTAMP    NOT NULL DEFAULT CURRENT_TIMESTAMP\n'
        ');',
        lang_label='SQL — log_sensor'
    )

    doc.add_paragraph(
        'Indeks UNIQUE pada pasangan (id_perangkat, timestamp) ditambahkan '
        'untuk menjamin idempotensi penyisipan ketika broker MQTT '
        'meretransmisi pesan akibat QoS 1. Tanpa indeks ini, satu publish '
        'yang sama bisa tersimpan dua kali. Dengan indeks ini, retransmisi '
        'akan jatuh ke klausul ON CONFLICT DO NOTHING di handler MQTT dan '
        'dibuang dengan aman.'
    )

    add_code_block(doc,
        'CREATE UNIQUE INDEX IF NOT EXISTS log_sensor_dedupe_idx\n'
        '    ON log_sensor (id_perangkat, timestamp);',
        lang_label='SQL — Indeks dedupe'
    )

    add_heading(doc, '2.2 Tabel pendaki', level=2)
    doc.add_paragraph(
        'Tabel pendaki menyimpan data registrasi yang dilakukan penjaga pos '
        'sebelum pendaki memulai pendakian. Kolom telepon_darurat dan '
        'tanggal_turun ditambahkan melalui migrasi inkremental ALTER TABLE '
        'IF NOT EXISTS, sehingga deployment lama yang sudah memiliki tabel '
        'pendaki versi awal tetap dapat di-upgrade tanpa kehilangan data.'
    )

    add_code_block(doc,
        'CREATE TABLE IF NOT EXISTS pendaki (\n'
        '    id              SERIAL      PRIMARY KEY,\n'
        '    nama_pendaki    VARCHAR(100) NOT NULL,\n'
        '    id_perangkat    VARCHAR(50)  NOT NULL,\n'
        '    telepon_darurat VARCHAR(30),\n'
        '    tanggal_naik    TIMESTAMP   NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
        '    tanggal_turun   TIMESTAMP,\n'
        '    status          VARCHAR(20) NOT NULL DEFAULT \'Mendaki\'\n'
        ');\n\n'
        'ALTER TABLE pendaki\n'
        '    ADD COLUMN IF NOT EXISTS telepon_darurat VARCHAR(30);\n'
        'ALTER TABLE pendaki\n'
        '    ADD COLUMN IF NOT EXISTS tanggal_turun TIMESTAMP;',
        lang_label='SQL — pendaki'
    )

    # ================================================================
    # Setup Container PostgreSQL
    # ================================================================
    add_heading(doc, '3. Setup Container PostgreSQL', level=1)
    doc.add_paragraph(
        'Basis data dijalankan sebagai container Docker yang dikonfigurasi '
        'melalui docker-compose.yml. Container tidak mempublish port 5432 '
        'ke host secara default karena seluruh akses dilakukan dari '
        'container backend melalui Docker network internal dengan hostname '
        '"postgres".'
    )

    add_code_block(doc,
        'services:\n'
        '  postgres:\n'
        '    image: postgres:15-alpine\n'
        '    container_name: altivex_postgres\n'
        '    environment:\n'
        '      POSTGRES_USER:     ${POSTGRES_USER:?belum diset di .env}\n'
        '      POSTGRES_PASSWORD: ${POSTGRES_PASSWORD:?belum diset di .env}\n'
        '      POSTGRES_DB:       ${POSTGRES_DB:?belum diset di .env}\n'
        '    volumes:\n'
        '      - pgdata:/var/lib/postgresql/data\n'
        '    restart: always\n'
        '    healthcheck:\n'
        '      test: ["CMD-SHELL", "pg_isready -U $${POSTGRES_USER}"]\n'
        '      interval: 10s\n'
        '      timeout:  5s\n'
        '      retries:  5',
        lang_label='YAML — docker-compose.yml'
    )

    doc.add_paragraph(
        'Konfigurasi healthcheck pg_isready memastikan container backend '
        'baru akan menjalin koneksi ke basis data setelah PostgreSQL siap '
        'menerima query. Hal ini menghindari race condition pada deployment '
        'pertama kali yang sebelumnya menyebabkan backend gagal membuat '
        'tabel.'
    )

    doc.add_paragraph('Connection string dibuat melalui template berikut, di mana username, password, dan nama database diisi otomatis oleh skrip bootstrap.sh saat deployment pertama:')

    add_code_block(doc,
        'DATABASE_URL=postgres://altivex_prod:<password-hex>@postgres:5432/altivex_db',
        lang_label='ENV — DATABASE_URL'
    )

    # ================================================================
    # Konfigurasi MQTT Mosquitto
    # ================================================================
    add_heading(doc, '4. Konfigurasi Broker MQTT (Mosquitto)', level=1)
    doc.add_paragraph(
        'Broker MQTT bertindak sebagai message bus antara device pendaki '
        'dan backend. Konfigurasi mosquitto.conf mengaktifkan otentikasi '
        'password dan menonaktifkan akses anonim. File password digenerate '
        'oleh skrip bootstrap.sh menggunakan utilitas mosquitto_passwd '
        'sehingga password yang sama di .env juga diketahui oleh broker.'
    )

    add_code_block(doc,
        'listener 1883 0.0.0.0\n'
        'allow_anonymous false\n'
        'password_file /mosquitto/config/passwd\n'
        'persistence true\n'
        'persistence_location /mosquitto/data/\n'
        'log_dest stdout\n'
        'log_type all',
        lang_label='mosquitto.conf'
    )

    doc.add_paragraph(
        'Topic yang digunakan ALTIVEX:'
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Topic'
    hdr[1].text = 'Arah'
    hdr[2].text = 'Pengirim — Penerima'
    for c in hdr:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(c, '1A1A1A')
        for p in c.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    rows = [
        ('altivex/sensor/data', 'Uplink', 'Device pendaki → Backend'),
        ('altivex/basecamp/cmd', 'Downlink', 'Backend → Device basecamp'),
        ('altivex/basecamp/ack', 'Uplink', 'Device basecamp → Backend'),
    ]
    for i, (topic, arah, ket) in enumerate(rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = topic
        cells[1].text = arah
        cells[2].text = ket
        # monospace for topic column
        for run in cells[0].paragraphs[0].runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)

    doc.add_paragraph()

    # ================================================================
    # Backend Startup Log
    # ================================================================
    add_heading(doc, '5. Log Startup Backend', level=1)
    doc.add_paragraph(
        'Berikut adalah keluaran log container backend pada deployment '
        'demo Situgede saat container baru diaktifkan. Empat baris kunci '
        'menunjukkan migrasi basis data, server HTTP, middleware otentikasi, '
        'dan subscription MQTT telah aktif.'
    )

    add_code_block(doc,
        '$ docker compose logs -f backend-demo\n'
        'altivex_backend_demo  | ✅ Database siap. Tabel log_sensor dan pendaki\n'
        '                      |    (dengan kolom telepon_darurat) tersedia.\n'
        'altivex_backend_demo  | ✅ Geofence ke-load dari "./frontend/GEO.json"\n'
        '                      |    (3 polygon segments).\n'
        'altivex_backend_demo  | 👀 Signal-lost watcher aktif\n'
        '                      |    (threshold=600s, interval=30s)\n'
        'altivex_backend_demo  | 🚀 Server ALTIVEX berjalan di http://0.0.0.0:8080\n'
        'altivex_backend_demo  | 🔐 AuthMiddleware aktif untuk endpoint mutating.\n'
        'altivex_backend_demo  | 🔑 Login basecamp aktif untuk user: demo\n'
        'altivex_backend_demo  | 🔌 Memulai Serial Reader di /dev/null...\n'
        'altivex_backend_demo  | 🛠️  Serial Writer task aktif (mpsc).\n'
        'altivex_backend_demo  | 📡 MQTT Subscriber aktif di topic:\n'
        '                      |    altivex/sensor/data (QoS=AtLeastOnce)',
        lang_label='Log — backend startup'
    )

    # ================================================================
    # Demo simulator publish
    # ================================================================
    add_heading(doc, '6. Pengujian Pengiriman Data dengan Simulator', level=1)
    doc.add_paragraph(
        'Untuk validasi pipeline tanpa hardware, disediakan simulator '
        'PowerShell dan Bash di scripts/demo-publisher.ps1 dan '
        'scripts/demo-publisher.sh. Simulator membangkitkan koordinat '
        'yang mengikuti loop bersepeda CIFOR–Situgede, lalu mempublish ke '
        'topic altivex/sensor/data dengan kredensial dari .env.demo.'
    )

    add_code_block(doc,
        'Indra@instance-20260424-035716:~/ALTIVEX$ ./scripts/demo-publisher.sh\n'
        '\n'
        '============================================================\n'
        'ALTIVEX Demo Publisher (bash)\n'
        '============================================================\n'
        'Device:    DEMO-CIFOR-01\n'
        'Broker:    altivex-demo.duckdns.org:1885\n'
        'Topic:     altivex/sensor/data\n'
        'Interval:  3 sec\n'
        'Loop:      10 minutes per round\n'
        'Mode:      LIVE PUBLISH\n'
        '============================================================\n'
        '\n'
        '[04:34:05] #1   loop=  0.0% bat=100% -> {"id_perangkat":"DEMO-CIFOR-01",\n'
        '                  "latitude":-6.554628,"longitude":106.751823,"battery":100}\n'
        '[04:34:09] #2   loop=  5.0% bat=100% -> {"id_perangkat":"DEMO-CIFOR-01",\n'
        '                  "latitude":-6.554186,"longitude":106.751243,"battery":100}\n'
        '[04:34:12] #3   loop= 10.1% bat=100% -> {"id_perangkat":"DEMO-CIFOR-01",\n'
        '                  "latitude":-6.553674,"longitude":106.750511,"battery":100}\n'
        '[04:34:15] #4   loop= 15.1% bat=100% -> {"id_perangkat":"DEMO-CIFOR-01",\n'
        '                  "latitude":-6.553111,"longitude":106.749623,"battery":100}',
        lang_label='Log — demo-publisher.sh'
    )

    # ================================================================
    # Backend log saat menerima publish
    # ================================================================
    add_heading(doc, '7. Log Backend Saat Menerima Publish', level=1)
    doc.add_paragraph(
        'Setelah simulator mempublish, backend mencatat tiga aktivitas '
        'utama untuk setiap pesan: (a) penerimaan publish dari broker, '
        '(b) penyisipan ke tabel log_sensor, dan (c) broadcast ke semua '
        'klien WebSocket dashboard yang terhubung.'
    )

    add_code_block(doc,
        '$ docker compose logs --tail=30 backend-demo\n'
        '\n'
        'altivex_backend_demo  | 📥 MQTT publish diterima:\n'
        '                      |    id=DEMO-CIFOR-01 lat=-6.554628 lon=106.751823\n'
        'altivex_backend_demo  | 💾 Insert OK ke log_sensor:\n'
        '                      |    id=DEMO-CIFOR-01 (1 row).\n'
        'altivex_backend_demo  | 📣 WS broadcast → 1 subscriber.\n'
        'altivex_backend_demo  | 📥 MQTT publish diterima:\n'
        '                      |    id=DEMO-CIFOR-01 lat=-6.554186 lon=106.751243\n'
        'altivex_backend_demo  | 💾 Insert OK ke log_sensor:\n'
        '                      |    id=DEMO-CIFOR-01 (1 row).\n'
        'altivex_backend_demo  | 📣 WS broadcast → 1 subscriber.\n'
        'altivex_backend_demo  | 📥 MQTT publish diterima:\n'
        '                      |    id=DEMO-CIFOR-01 lat=-6.553674 lon=106.750511\n'
        'altivex_backend_demo  | 💾 Insert OK ke log_sensor:\n'
        '                      |    id=DEMO-CIFOR-01 (1 row).\n'
        'altivex_backend_demo  | 📣 WS broadcast → 1 subscriber.',
        lang_label='Log — backend menerima MQTT'
    )

    doc.add_paragraph(
        'Tag emoji digunakan secara konsisten di seluruh log backend untuk '
        'mempermudah grep dan debugging:'
    )

    table = doc.add_table(rows=6, cols=2)
    table.style = 'Light Grid Accent 1'
    hdr = table.rows[0].cells
    hdr[0].text = 'Tag'
    hdr[1].text = 'Arti'
    for c in hdr:
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
        shade_cell(c, '1A1A1A')
        for p in c.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    tag_rows = [
        ('📥', 'Publish MQTT diterima dari pendaki'),
        ('💾', 'Insert ke tabel log_sensor sukses'),
        ('↩️',  'Dedupe (publish duplikat dari retransmit broker)'),
        ('📣', 'Broadcast WebSocket ke dashboard'),
        ('🚨', 'Alert otomatis di-publish ke basecamp'),
    ]
    for i, (tag, ket) in enumerate(tag_rows, start=1):
        cells = table.rows[i].cells
        cells[0].text = tag
        cells[1].text = ket
        for run in cells[0].paragraphs[0].runs:
            run.font.size = Pt(14)

    doc.add_paragraph()

    # ================================================================
    # Verifikasi data di basis data
    # ================================================================
    add_heading(doc, '8. Verifikasi Data di Basis Data', level=1)
    doc.add_paragraph(
        'Untuk memastikan koordinat benar-benar tersimpan, dapat dilakukan '
        'query langsung ke container PostgreSQL dengan perintah berikut. '
        'Hasil query menampilkan baris-baris terbaru pada tabel log_sensor '
        'beserta timestamp aktual saat penyisipan.'
    )

    add_code_block(doc,
        '$ docker compose exec postgres-demo psql \\\n'
        '    -U altivex_demo -d altivex_demo_db \\\n'
        '    -c "SELECT id_perangkat, latitude, longitude, battery, timestamp \\\n'
        '        FROM log_sensor ORDER BY timestamp DESC LIMIT 5;"',
        lang_label='Shell — query log_sensor'
    )

    add_code_block(doc,
        ' id_perangkat  |  latitude  | longitude  | battery |       timestamp\n'
        '---------------+------------+------------+---------+----------------------\n'
        ' DEMO-CIFOR-01 | -6.553674  | 106.750511 |     100 | 2026-05-21 04:34:12\n'
        ' DEMO-CIFOR-01 | -6.554186  | 106.751243 |     100 | 2026-05-21 04:34:09\n'
        ' DEMO-CIFOR-01 | -6.554628  | 106.751823 |     100 | 2026-05-21 04:34:05\n'
        '(3 rows)',
        lang_label='Output — psql'
    )

    doc.add_paragraph(
        'Hasil query memperlihatkan tiga baris yang konsisten dengan tiga '
        'pesan publish pertama dari simulator. Latitude dan longitude '
        'tersimpan dengan presisi 6 desimal (±0,1 meter di equator), '
        'memenuhi kebutuhan akurasi tracking pendaki.'
    )

    # ================================================================
    # Penutup
    # ================================================================
    add_heading(doc, '9. Penutup', level=1)
    doc.add_paragraph(
        'Dokumentasi ini menunjukkan bahwa pipeline data ALTIVEX dari '
        'device hingga basis data berjalan utuh: broker MQTT menerima '
        'publish, backend mem-parsing payload, menyimpan ke tabel log_sensor '
        'dengan jaminan idempotensi, dan menyebarkan posisi ke dashboard '
        'melalui WebSocket dalam waktu kurang dari satu detik. Validasi ini '
        'menjadi dasar uji integrasi bahwa sistem siap menerima data dari '
        'hardware pendaki sesungguhnya pada tahap pengujian lapangan.'
    )

    # ================================================================
    # Save
    # ================================================================
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    repo_root = Path(__file__).resolve().parent.parent
    out_dir = repo_root / 'docs'
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / 'DOKUMENTASI_DATABASE_DAN_LOGS.docx'
    build_document(str(out_path))
