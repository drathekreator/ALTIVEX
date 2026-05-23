"""
Generate dokumen DOCX untuk BAB 3 TAHAP PELAKSANAAN proyek ALTIVEX.

Sub-bab:
  3.1 Implementasi Realtime Database
  3.2 Implementasi Web

Output: docs/BAB_3_TAHAP_PELAKSANAAN.docx
"""

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# --------------------------------------------------------------------
# Style helpers (re-used dari generate-db-docs.py)
# --------------------------------------------------------------------
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def add_code_block(doc, text, lang_label=None):
    if lang_label:
        p = doc.add_paragraph()
        run = p.add_run(f'[{lang_label}]')
        run.italic = True
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)

    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F4F4F4')
    p_pr.append(shd)

    pBdr = OxmlElement('w:pBdr')
    for edge in ('top', 'bottom', 'left', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), 'CCCCCC')
        pBdr.append(b)
    p_pr.append(pBdr)

    run = p.add_run(text)
    run.font.name = 'Consolas'
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:ascii'), 'Consolas')
    rFonts.set(qn('w:hAnsi'), 'Consolas')
    run.font.size = Pt(9)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return h


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)


def header_table(table, hex_bg='1A1A1A'):
    for c in table.rows[0].cells:
        shade_cell(c, hex_bg)
        for p in c.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


# --------------------------------------------------------------------
# Build document
# --------------------------------------------------------------------
def build_document(out_path):
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

    # Default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ================================================================
    # Title
    # ================================================================
    h = doc.add_heading('BAB 3\nTAHAP PELAKSANAAN', level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ================================================================
    # 3.1 Implementasi Realtime Database
    # ================================================================
    add_heading(doc, '3.1 Implementasi Realtime Database', level=1)

    doc.add_paragraph(
        'Tahap implementasi realtime database pada sistem ALTIVEX dimulai '
        'dengan pemilihan kombinasi PostgreSQL sebagai sistem manajemen '
        'basis data dan Eclipse Mosquitto sebagai message broker MQTT '
        'untuk transportasi data telemetri dari device pendaki menuju '
        'backend. Pemilihan PostgreSQL dilatari oleh kemampuan ACID, '
        'dukungan terhadap indeks unik komposit, dan ekosistem yang matang '
        'untuk aplikasi yang membutuhkan integritas data tinggi seperti '
        'pelacakan posisi pendaki secara berkala. Mosquitto dipilih '
        'karena ringan, mendukung otentikasi password, dan menjadi '
        'implementasi referensi protokol MQTT 3.1.1.'
    )

    doc.add_paragraph(
        'Skema basis data dirancang sederhana namun cukup untuk kebutuhan '
        'sistem, terdiri atas dua tabel utama, yaitu log_sensor dan '
        'pendaki. Tabel log_sensor menyimpan setiap titik koordinat yang '
        'dikirim device pendaki, sedangkan tabel pendaki menyimpan data '
        'registrasi pendaki sebelum melakukan perjalanan. Skema kedua '
        'tabel beserta indeks dedupe ditampilkan pada blok kode berikut.'
    )

    add_code_block(doc,
        'CREATE TABLE IF NOT EXISTS log_sensor (\n'
        '    id            SERIAL PRIMARY KEY,\n'
        '    id_perangkat  VARCHAR(50) NOT NULL,\n'
        '    latitude      DOUBLE PRECISION NOT NULL,\n'
        '    longitude     DOUBLE PRECISION NOT NULL,\n'
        '    battery       SMALLINT,\n'
        '    timestamp     TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP\n'
        ');\n\n'
        'CREATE UNIQUE INDEX IF NOT EXISTS log_sensor_dedupe_idx\n'
        '    ON log_sensor (id_perangkat, timestamp);\n\n'
        'CREATE TABLE IF NOT EXISTS pendaki (\n'
        '    id              SERIAL PRIMARY KEY,\n'
        '    nama_pendaki    VARCHAR(100) NOT NULL,\n'
        '    id_perangkat    VARCHAR(50)  NOT NULL,\n'
        '    telepon_darurat VARCHAR(30),\n'
        '    tanggal_naik    TIMESTAMP NOT NULL DEFAULT CURRENT_TIMESTAMP,\n'
        '    tanggal_turun   TIMESTAMP,\n'
        '    status          VARCHAR(20) NOT NULL DEFAULT \'Mendaki\'\n'
        ');',
        lang_label='SQL — Skema basis data ALTIVEX'
    )
    add_caption(doc, 'Listing 3.1. Skema basis data PostgreSQL sistem ALTIVEX.')

    doc.add_paragraph(
        'Migrasi skema dijalankan secara otomatis oleh backend pada saat '
        'startup melalui perintah CREATE TABLE IF NOT EXISTS dan '
        'ALTER TABLE IF NOT EXISTS untuk kolom inkremental seperti '
        'telepon_darurat dan tanggal_turun. Pendekatan ini memungkinkan '
        'sistem dideploy ulang berkali-kali tanpa langkah migrasi manual, '
        'sekaligus mempertahankan kompatibilitas dengan basis data lama '
        'yang sudah berisi data pendaki sebelumnya.'
    )

    doc.add_paragraph(
        'Untuk menjamin sifat realtime, sistem menggunakan pola publish '
        'and subscribe MQTT. Device pendaki mempublish payload JSON ke '
        'topic altivex/sensor/data setiap lima detik berisi id_perangkat, '
        'latitude, longitude, dan persen baterai. Backend yang berperan '
        'sebagai subscriber menerima payload, melakukan validasi koordinat, '
        'menyisipkan ke tabel log_sensor, dan menyebarkan kembali ke '
        'seluruh dashboard penjaga pos melalui koneksi WebSocket. Daftar '
        'topic MQTT yang digunakan ditampilkan pada Tabel 3.1.'
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Topic'
    table.rows[0].cells[1].text = 'Arah'
    table.rows[0].cells[2].text = 'Pengirim → Penerima'
    header_table(table)

    rows = [
        ('altivex/sensor/data',  'Uplink',   'Device pendaki → Backend'),
        ('altivex/basecamp/cmd', 'Downlink', 'Backend → Device basecamp'),
        ('altivex/basecamp/ack', 'Uplink',   'Device basecamp → Backend'),
    ]
    for i, (a, b, c) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = a
        table.rows[i].cells[1].text = b
        table.rows[i].cells[2].text = c
        for run in table.rows[i].cells[0].paragraphs[0].runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
    add_caption(doc, 'Tabel 3.1. Daftar topic MQTT ALTIVEX dan arah komunikasinya.')
    doc.add_paragraph()

    doc.add_paragraph(
        'Tantangan utama dalam pipeline realtime adalah kemungkinan '
        'duplikasi pesan akibat retransmisi broker pada level QoS 1. '
        'Tanpa mekanisme dedupe, satu publish yang sama dapat tersimpan '
        'dua kali dan menimbulkan polyline yang ganda di peta. Sistem '
        'mengatasi hal ini dengan dua lapis pertahanan, yaitu indeks '
        'unik komposit log_sensor_dedupe_idx pada pasangan '
        '(id_perangkat, timestamp) dan klausul ON CONFLICT DO NOTHING '
        'pada perintah penyisipan. Kombinasi ini menjadikan operasi '
        'INSERT idempoten sehingga retransmisi diserap dengan aman.'
    )

    doc.add_paragraph(
        'Container PostgreSQL dijalankan melalui Docker Compose dengan '
        'volume persisten pgdata yang ditempatkan terpisah dari kontainer '
        'sehingga data tidak hilang saat container di-rebuild. Healthcheck '
        'pg_isready dikonfigurasi pada interval sepuluh detik dengan lima '
        'kali percobaan ulang. Hasilnya, container backend baru memulai '
        'koneksi setelah PostgreSQL benar-benar siap menerima query, '
        'menghindari race condition pada deployment pertama kali.'
    )

    add_code_block(doc,
        'services:\n'
        '  postgres:\n'
        '    image: postgres:15-alpine\n'
        '    environment:\n'
        '      POSTGRES_USER:     ${POSTGRES_USER}\n'
        '      POSTGRES_PASSWORD: ${POSTGRES_PASSWORD}\n'
        '      POSTGRES_DB:       ${POSTGRES_DB}\n'
        '    volumes:\n'
        '      - pgdata:/var/lib/postgresql/data\n'
        '    healthcheck:\n'
        '      test: ["CMD-SHELL", "pg_isready -U $${POSTGRES_USER}"]\n'
        '      interval: 10s\n'
        '      retries:  5',
        lang_label='YAML — docker-compose.yml'
    )
    add_caption(doc, 'Listing 3.2. Definisi container PostgreSQL pada Docker Compose.')

    doc.add_paragraph(
        'Pengujian integrasi pipeline dilakukan dengan mempublish payload '
        'tiruan menggunakan utilitas mosquitto_pub dari laptop. Backend '
        'mencatat alur pemrosesan tiap pesan ke standard output container '
        'dengan tag emoji yang konsisten guna memudahkan grep dan analisis. '
        'Kutipan log saat backend menerima publish ditampilkan pada '
        'Listing 3.3.'
    )

    add_code_block(doc,
        'altivex_backend  | 📥 MQTT publish diterima:\n'
        '                 |    id=DEMO-CIFOR-01 lat=-6.554628 lon=106.751823\n'
        'altivex_backend  | 💾 Insert OK ke log_sensor:\n'
        '                 |    id=DEMO-CIFOR-01 (1 row).\n'
        'altivex_backend  | 📣 WS broadcast → 1 subscriber.',
        lang_label='Log — Backend menerima publish MQTT'
    )
    add_caption(doc, 'Listing 3.3. Kutipan log backend ketika menerima dan memproses publish.')

    doc.add_paragraph(
        'Arti tag emoji pada log adalah sebagai berikut: 📥 menandakan '
        'pesan publish telah diterima dari broker, 💾 menyatakan baris '
        'baru berhasil disisipkan ke tabel log_sensor, ↩️ menandakan '
        'pesan duplikat yang diserap oleh klausul ON CONFLICT, dan 📣 '
        'menyatakan posisi telah di-broadcast ke dashboard melalui '
        'WebSocket. Verifikasi akhir bahwa data benar-benar tersimpan '
        'dilakukan dengan query langsung ke basis data sebagaimana '
        'ditunjukkan pada Listing 3.4.'
    )

    add_code_block(doc,
        '$ docker compose exec postgres psql -U altivex_prod -d altivex_db \\\n'
        '    -c "SELECT id_perangkat, latitude, longitude, battery, timestamp \\\n'
        '        FROM log_sensor ORDER BY timestamp DESC LIMIT 3;"\n\n'
        ' id_perangkat  |  latitude  | longitude  | battery |       timestamp\n'
        '---------------+------------+------------+---------+----------------------\n'
        ' DEMO-CIFOR-01 | -6.553674  | 106.750511 |     100 | 2026-05-21 04:34:12\n'
        ' DEMO-CIFOR-01 | -6.554186  | 106.751243 |     100 | 2026-05-21 04:34:09\n'
        ' DEMO-CIFOR-01 | -6.554628  | 106.751823 |     100 | 2026-05-21 04:34:05\n'
        '(3 rows)',
        lang_label='Shell + Output — Query verifikasi'
    )
    add_caption(doc, 'Listing 3.4. Verifikasi keberadaan data sensor melalui query langsung psql.')

    doc.add_paragraph(
        'Berdasarkan hasil pengujian tersebut, jalur data dari device '
        'pendaki hingga basis data tervalidasi berfungsi utuh dengan '
        'jaminan idempotensi dan ketahanan terhadap retransmisi. Latensi '
        'rata-rata dari publish ke broadcast WebSocket terukur di bawah '
        'satu detik, memenuhi karakteristik realtime untuk pemantauan '
        'pendaki.'
    )

    # Page break before next section
    doc.add_page_break()

    # ================================================================
    # 3.2 Implementasi Web
    # ================================================================
    add_heading(doc, '3.2 Implementasi Web', level=1)

    doc.add_paragraph(
        'Antarmuka web ALTIVEX dirancang sebagai dashboard tunggal yang '
        'dijalankan di browser penjaga pos pendakian. Frontend dibangun '
        'menggunakan kombinasi HTML5, JavaScript vanilla, pustaka Leaflet.js '
        'untuk peta interaktif, dan pustaka Turf.js untuk komputasi '
        'geospasial sisi klien. Pendekatan vanilla tanpa framework '
        'berorientasi komponen seperti React atau Vue dipilih agar '
        'ukuran initial load tetap kecil dan tidak menambah ketergantungan '
        'build pipeline yang tidak diperlukan untuk skala dashboard '
        'tunggal.'
    )

    doc.add_paragraph(
        'Sisi server diimplementasikan dengan bahasa Rust dan framework '
        'Actix-web. Pemilihan Rust didasari kebutuhan performa tinggi '
        'pada pemrosesan publish MQTT yang berfrekuensi padat dan '
        'jaminan memory safety yang ketat tanpa overhead garbage '
        'collector. Backend bertindak sekaligus sebagai server REST '
        'API, pengelola koneksi WebSocket, klien MQTT subscriber, '
        'evaluator geofence, dan penerbit perintah ke device basecamp. '
        'Static asset dashboard juga di-serve langsung oleh backend '
        'menggunakan modul actix-files sehingga tidak perlu web server '
        'tambahan.'
    )

    doc.add_paragraph(
        'Arsitektur tiga lapisan sistem ALTIVEX dirangkum pada '
        'Tabel 3.2 yang memetakan teknologi dengan peran tiap lapisan.'
    )

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Lapisan'
    table.rows[0].cells[1].text = 'Teknologi'
    table.rows[0].cells[2].text = 'Tugas'
    header_table(table)

    rows = [
        ('Edge / Device',
         'ESP32, GPS NEO-6M, modul LoRa atau Wi-Fi, PubSubClient',
         'Mengakuisisi koordinat dan mempublish ke broker MQTT.'),
        ('Cloud Backend',
         'Rust + Actix-web, PostgreSQL 15, Mosquitto 2, Docker Compose, nginx + certbot',
         'Mengelola REST API, WebSocket, MQTT subscriber, geofence engine, dan penyimpanan basis data.'),
        ('Web Dashboard',
         'HTML5 + JavaScript vanilla, Leaflet.js, Turf.js, SheetJS',
         'Menampilkan peta real-time, alert, riwayat, dan ekspor data Excel.'),
    ]
    for i, (lap, tek, tug) in enumerate(rows, start=1):
        table.rows[i].cells[0].text = lap
        table.rows[i].cells[1].text = tek
        table.rows[i].cells[2].text = tug
    add_caption(doc, 'Tabel 3.2. Pemetaan teknologi pada arsitektur tiga lapisan ALTIVEX.')
    doc.add_paragraph()

    doc.add_paragraph(
        'Dashboard memiliki dua tab utama. Tab pertama adalah Peta Live '
        'yang menampilkan posisi seluruh pendaki yang sedang aktif '
        'beserta polyline jalur resmi pendakian, area buffer geofence, '
        'dan sidebar berisi kartu pendaki yang dikelompokkan berdasarkan '
        'status aman atau di luar koridor. Tab kedua adalah Kelola '
        'Pendaki yang berisi formulir registrasi, tabel riwayat semua '
        'pendaki yang pernah tercatat, serta tombol penyelesaian '
        'pendakian dan ekspor riwayat ke berkas Excel. Pemisahan tab '
        'memudahkan penjaga pos berfokus pada pemantauan tanpa '
        'terganggu daftar administratif.'
    )

    doc.add_paragraph(
        'Komunikasi real-time antara backend dan dashboard menggunakan '
        'WebSocket pada path /ws. Saat dashboard pertama kali dibuka, '
        'klien melakukan upgrade koneksi HTTP ke WebSocket dan '
        'mendaftar sebagai subscriber broadcast channel di backend. '
        'Setiap publish posisi yang masuk dari MQTT akan didistribusikan '
        'ke semua dashboard aktif dalam bentuk pesan JSON. Untuk '
        'menjaga ketahanan saat WebSocket terputus, dashboard juga '
        'memiliki polling fallback ke endpoint /api/sensor/latest '
        'setiap tiga puluh detik sehingga peta tetap mendapat '
        'pembaruan posisi.'
    )

    add_code_block(doc,
        'const ws = new WebSocket(\n'
        '    (location.protocol === "https:" ? "wss://" : "ws://") +\n'
        '    location.host + "/ws"\n'
        ');\n\n'
        'ws.onmessage = (event) => {\n'
        '    const data = JSON.parse(event.data);\n'
        '    latestDataPerDevice[data.id_perangkat] = data;\n'
        '    renderHikerCards();\n'
        '    updateMarker(data);\n'
        '};\n\n'
        'ws.onclose = () => setTimeout(reconnectWs, 2000);',
        lang_label='JavaScript — Inisialisasi WebSocket'
    )
    add_caption(doc, 'Listing 3.5. Inisialisasi koneksi WebSocket dari dashboard.')

    doc.add_paragraph(
        'Otentikasi dashboard dilakukan dengan dua tahap. Tahap pertama '
        'adalah login yang memverifikasi BASECAMP_USERNAME dan '
        'BASECAMP_PASSWORD terhadap nilai yang disimpan di environment '
        'backend. Verifikasi menggunakan perbandingan waktu konstan untuk '
        'menutup celah timing attack. Setelah login berhasil, backend '
        'mengembalikan API_AUTH_TOKEN yang disimpan oleh dashboard di '
        'localStorage. Token ini kemudian disertakan sebagai header '
        'Authorization Bearer pada setiap permintaan REST API yang '
        'bersifat mutating, seperti registrasi pendaki dan penyelesaian '
        'pendakian. Endpoint publik seperti /api/sensor/latest dan /ws '
        'tidak memerlukan token untuk menjaga sederhananya akses '
        'dashboard.'
    )

    doc.add_paragraph(
        'Geofence di-evaluasi dua kali untuk memberikan pengalaman '
        'pengguna yang responsif sekaligus jaminan deteksi yang ketat. '
        'Evaluasi pertama dilakukan di sisi klien menggunakan Turf.js '
        'sehingga banner peringatan langsung muncul tanpa menunggu '
        'round-trip ke backend. Evaluasi kedua dilakukan di backend '
        'Rust dengan crate geo dan geojson, mengeluarkan perintah '
        'otomatis ke device basecamp via topic altivex/basecamp/cmd '
        'tanpa intervensi penjaga. Pendekatan ganda ini memastikan '
        'alert tetap berfungsi meskipun browser dashboard sedang tidak '
        'aktif.'
    )

    doc.add_paragraph(
        'Beberapa fitur tambahan diimplementasikan untuk meningkatkan '
        'kenyamanan operasional, di antaranya mode gelap dengan '
        'penyesuaian kontras otomatis, ekspor riwayat pendakian ke '
        'berkas Excel berbasis SheetJS, indikator persen baterai per '
        'device beserta notifikasi browser ketika baterai turun di '
        'bawah lima belas persen, dan banner peringatan in-app yang '
        'dapat di-klik untuk melompat ke kartu pendaki yang berada di '
        'luar koridor. Daftar fitur dashboard yang telah '
        'diimplementasikan dirangkum pada Tabel 3.3.'
    )

    table = doc.add_table(rows=10, cols=2)
    table.style = 'Light Grid Accent 1'
    table.rows[0].cells[0].text = 'Fitur'
    table.rows[0].cells[1].text = 'Deskripsi Singkat'
    header_table(table)

    fitur_rows = [
        ('Peta Live',
         'Marker bergerak otomatis sesuai publish posisi, polyline jalur, dan area buffer geofence.'),
        ('Sidebar Alert',
         'Kartu pendaki di luar koridor terangkat ke atas dengan warna kontras.'),
        ('Banner Peringatan',
         'Notifikasi in-app yang dapat di-klik untuk fokus ke pendaki bermasalah.'),
        ('Indikator Baterai',
         'Persen baterai tiap device ditampilkan dengan warna gradasi merah hingga hijau.'),
        ('Riwayat Polyline',
         'Tombol PATH menampilkan jejak pendakian untuk pendaki tertentu.'),
        ('Tabel Kelola Pendaki',
         'Daftar registrasi dan riwayat dengan aksi edit, hapus, dan selesai.'),
        ('Ekspor Excel',
         'Mengunduh seluruh riwayat ke berkas .xlsx dengan format siap audit.'),
        ('Mode Gelap',
         'Toggle tema gelap dengan deteksi preferensi sistem operasi.'),
        ('Login Basecamp',
         'Otentikasi tunggal dengan token Bearer pada permintaan mutating.'),
    ]
    for i, (fitur, desk) in enumerate(fitur_rows, start=1):
        table.rows[i].cells[0].text = fitur
        table.rows[i].cells[1].text = desk
    add_caption(doc, 'Tabel 3.3. Fitur utama dashboard ALTIVEX.')
    doc.add_paragraph()

    doc.add_paragraph(
        'Deployment dashboard dilakukan bersamaan dengan backend dalam '
        'satu container Docker. Reverse proxy nginx pada host VM '
        'menerima permintaan HTTPS pada domain '
        'altivex-pangrango.duckdns.org untuk produksi dan '
        'altivex-demo.duckdns.org untuk pengujian, kemudian meneruskan '
        'ke backend di port internal. Sertifikat TLS diperoleh secara '
        'otomatis dari Let\'s Encrypt menggunakan utilitas certbot '
        'dengan jadwal pembaruan bawaan setiap enam puluh hari, '
        'sehingga komunikasi antara dashboard dan backend terenkripsi '
        'sepanjang waktu tanpa intervensi manual.'
    )

    doc.add_paragraph(
        'Pengujian end-to-end dashboard dilakukan dengan skenario '
        'simulasi loop bersepeda di kawasan CIFOR–Situgede, Bogor, '
        'menggunakan device demo yang mempublish posisi setiap tiga '
        'detik. Hasil pengujian menunjukkan marker pendaki bergerak '
        'mengikuti rute yang dirancang dengan halus, banner peringatan '
        'muncul tepat ketika koordinat keluar dari koridor, dan '
        'indikator baterai turun perlahan sesuai pola yang '
        'dikonfigurasi pada device. Hasil ini memvalidasi bahwa '
        'integrasi antara backend, basis data realtime, dan dashboard '
        'web telah berfungsi sesuai spesifikasi.'
    )

    # ================================================================
    # Save
    # ================================================================
    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == '__main__':
    repo_root = Path(__file__).resolve().parent.parent
    out_dir = repo_root / 'docs'
    out_dir.mkdir(exist_ok=True)
    out_path = out_dir / 'BAB_3_TAHAP_PELAKSANAAN.docx'
    build_document(str(out_path))
